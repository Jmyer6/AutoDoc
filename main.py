from PySide6.QtGui import QCloseEvent,QColor
from PySide6.QtWidgets import QApplication,QVBoxLayout,QHBoxLayout,QWidget,QFrame,QSizePolicy
from PySide6.QtCore import QThread,Signal,QSize,Qt

import sys,os
from widget import ButtonNormal,WindowsFramelessWindow,LineEdit,LabelNormal,TextEdit,\
    SpinBox,Icon,CarouselFrame,IconView,InfoType,CircularProgressBar,ComBoBox,CheckRoundBox,\
    Message,NameRuleList,TaskList
from qss import QssPlusClass,Color_,ColorStyle,BorderStyle,Font_,FontSize,FontWeight,TextAlign
from lib import Position

from docx import Document
from openpyxl import load_workbook

loadWordIndex=0
loadKeyIndex=1
loadExeclIndex=2
loadDataIndex=3
loadNameIndex=4
loadRunIndex=5
loadWaitIndex=6

textwidth=113

class MainWidget(WindowsFramelessWindow):
    def __init__(self, parent=None):
        super().__init__(parent) 

        self._docLink=None
        self._docLoadThread=None
        self._docData=None
        self._docKey=None
        self._docKeyNumber=None
        self._docKeyThread=None
        self._docKeyCache=None

        self._execlLink=None
        self._execlLoadThread=None
        self._execlData=None
        self._execlDataIndex=None

        self._nameList=[]

        self.setWindowTitle("Word模板批量修改")
        self.setWindowIcon(Icon.word2)

        self.__SetLayout()
        self.__Resize()
        self.sizechanged.connect(self.__Resize)

    def __SetLayout(self):
        self._frame=CarouselFrame(Position.LEFT,parent=self.userFrame)

        self._loadWord=LoadWord()
        self._loadWord.nextpressed.connect(self.__WordNext)

        self._loadWait=WaitWidget()
        self._frame._moveframe1._animation._animation.finished.connect(self.__LoadWaitStop)

        self._loadKey=LoadKey()
        self._loadKey.beforpressed.connect(self._frame.Go)
        self._loadKey.checkpressed.connect(self.__KeyCheckNumber)
        self._loadKey.nextpressed.connect(self.__KeyNext)

        self._loadExecl=LoadExecl()
        self._loadExecl.beforpressed.connect(self._frame.Go)
        self._loadExecl.nextpressed.connect(self.__ExeclNext)

        self._loadData=LoadData()
        self._loadData.beforpressed.connect(self._frame.Go)
        self._loadData.nextpressed.connect(self.__DateNext)

        self._loadName=LoadName()
        self._loadName.beforpressed.connect(self._frame.Go)
        self._loadName.nextpressed.connect(self.__NameNext)

        self._savedoc=SaveDoc()
        self._savedoc.beforpressed.connect(self._frame.Go)
        self._savedoc.firsted.connect(self._frame.Go)
        self._savedoc.openfiled.connect(self.__OpenFile)
        self._savedoc.runed.connect(self.__Run)

        self._frame.AddItem(self._loadWord)
        self._frame.AddItem(self._loadKey)
        self._frame.AddItem(self._loadExecl)
        self._frame.AddItem(self._loadData)
        self._frame.AddItem(self._loadName)
        self._frame.AddItem(self._savedoc)

        self._frame.AddItem(self._loadWait)

        self._frame.Go(loadWordIndex)
    
    def __LoadWaitStop(self):
        if self._frame._index!=loadWaitIndex:
            self._loadWait._progress.Stop()

        
    def __Resize(self):
        self._frame.move(0,0)
        self._frame.setFixedSize(self.userFrame.size())

        for item in self._frame._list:
            item.setFixedSize(self.userFrame.size())

    def showEvent(self, event) -> None:
        self.__Resize()
        return super().showEvent(event)
    
    def __WordNext(self,link:str):
        if self._loadWord.IsWord(link):
            self._docLink=link

            self._loadWait.SetText("Word文件加载中...(文件不能使用office、wps等打开)")
            self._loadWait._progress.Start()
            self._frame.Go(loadWaitIndex)

            self._docLoadThread=LoadWordThread(self._docLink)
            self._docLoadThread.loaded.connect(self.__WordSet)
            self._docLoadThread.errored.connect(self.__WordError)
            self._docLoadThread.start()
        else:
            self.window()._info.ShowInfo(InfoType.WARNING,"提示:","载入文件不是Word文件。")

    def __WordSet(self,doc):
        self._docData=doc
        self._frame.Go(loadKeyIndex,True)

        self._loadKey._wordlink.setText(self._docLink)
        self.window()._info.ShowInfo(InfoType.SUCCESS,"提示:","Word模板文件载入成功。")

    def __WordError(self):
        self._frame.Go(loadWordIndex)
        self.window()._info.ShowInfo(InfoType.ERROR,"提示:","Word模板文件无法加载,请检测路径和文件是否损坏或已被打开。")


    def __KeyNext(self,key:str):
        if key=="":
            self.window()._info.ShowInfo(InfoType.WARNING,"提示:","关键词不能为空。")
            return
        
        elif self._docKeyCache is not None and self._docKeyNumber is not None and self._docKeyCache==key and self._docKeyNumber!=0:
            self._docKey=key
            self._frame.Go(loadExeclIndex)
            return
        else:
            self._loadWait.SetText("Word关键词检测中...")
            self._loadWait._progress.Start()
            self._frame.Go(loadWaitIndex)

            self._docKey=key
            self._docKeyThread=KeyThread(key,self._docData)
            self._docKeyThread.checked.connect(self.__KeyNumberSet)
            self._docKeyThread.start()

    def __KeyNumberSet(self,Number:int):
        self._docKeyNumber=Number

        if Number==0:
            self.window()._info.ShowInfo(InfoType.ERROR,"提示:","Word模板文件中含关键词["+str(self._docKey)+"]有0个。")  
            self._frame.Go(loadKeyIndex)
        else:
            self.window()._info.ShowInfo(InfoType.SUCCESS,"提示:","Word模板文件中含关键词["+str(self._docKey)+"]有"+str(Number)+"个。")
            self._frame.Go(loadExeclIndex,True)

    def __KeyCheckNumber(self,key:str):
        if self._docKeyCache is not None and self._docKeyNumber is not None and self._docKeyCache==key and self._docKeyNumber != 0:
            self.window()._info.ShowInfo(InfoType.SUCCESS,"提示:","Word模板文件中含关键词["+str(self._docKeyCache)+"]有"+str(self._docKeyNumber)+"个。")
            return
        
        self._docKeyCache=key
        self._docKeyThread=KeyThread(key,self._docData)
        self._docKeyThread.checked.connect(self.__KeyShowNumber)
        self._docKeyThread.start()

    def __KeyShowNumber(self,Number:int):
        self._docKeyNumber=Number
        if self._docKeyNumber==0:
            self.window()._info.ShowInfo(InfoType.ERROR,"提示:","Word模板文件中含关键词["+str(self._docKeyCache)+"]有0个。")  
        else:
            self.window()._info.ShowInfo(InfoType.SUCCESS,"提示:","Word模板文件中含关键词["+str(self._docKeyCache)+"]有"+str(Number)+"个。")
        

    def __ExeclNext(self,link:str):
        if self._loadExecl.IsExecl(link):
            self._execlLink=link

            self._loadWait.SetText("Execl文件加载中...(文件不能使用office、wps等打开)")
            self._loadWait._progress.Start()
            self._frame.Go(loadWaitIndex)

            self._execlLoadThread=LoadExeclThread(self._execlLink)
            self._execlLoadThread.loaded.connect(self.__ExeclSet)
            self._execlLoadThread.errored.connect(self.__ExeclError)
            self._execlLoadThread.start()
        else:
            self.window()._info.ShowInfo(InfoType.WARNING,"提示:","载入文件不是Execl文件。")

    def __ExeclSet(self,execl):
        self._execlData=execl
        self._frame.Go(loadDataIndex,True)

        self._loadData._execldatapage.Clear()
        self._loadData._execldatapage.Addtexts(self._execlData.sheetnames)
        self._loadData._execldatapage.SetCheck(0)

        self._loadData._execllink.setText(self._execlLink)
        self.window()._info.ShowInfo(InfoType.SUCCESS,"提示:","Execl文件载入成功。")

    def __ExeclError(self):
        self._frame.Go(loadExeclIndex)
        self.window()._info.ShowInfo(InfoType.ERROR,"提示:","Execl文件无法加载,请检测路径和文件是否损坏或已被打开。")


    def __DateNext(self,position:int):

        sheet = self._execlData.worksheets[position]
        _maxrow=sheet.max_row
        if _maxrow!=self._docKeyNumber:
            self.window()._info.ShowInfo(InfoType.ERROR,"提示:","Execl数据表中有"+str(_maxrow)+"行数据,Word模板关键词["+str(self._docKey)+"]有"+str(self._docKeyNumber)+"个,数量不一致,请检查。")
            return

        _maxcol=sheet.max_column
        for _x in range(1,_maxcol+1):
            for _y in range(1,_maxrow+1):
                if sheet.cell(row=_y, column=_x).value==None:
                    self.window()._info.ShowInfo(InfoType.ERROR,"提示:","Execl数据表中第"+str(_x)+"行第"+str(_y)+"列中存在空值,数据不能为空,请检查。")
                    return 
                
        self._execlDataIndex=position

        _path,_fullname = os.path.split(self._docLink)
        _name,_suffix = os.path.splitext(_fullname)

        self._loadName._savelink.setText(_path)
        self._loadName._nameRule.AddRightButton(_suffix)
        if self._loadName._nameRule._rulelist==[]:
            self._loadName._nameRule.AddNameRule("文本",_name)
            self._loadName._nameRule.AddNameRule("数字",1)
        self._loadName._nameRule.AddExeclData(self._execlData)
        self._loadName._nameRule.AddExeclDataIndex(self._execlDataIndex)

        self._frame.Go(loadNameIndex)

    def __NameNext(self,RuleList:list):

        _number=self._execlData.worksheets[self._execlDataIndex].max_column

        _filelink=self._loadName._savelink.text()

        _path,_fullname = os.path.split(self._docLink)
        _name,_suffix = os.path.splitext(_fullname)

        self._savedoc._savelist.Clear()

        for i in range(_number):
            _filename=""
            for _rule in RuleList:
                if _rule[1]=="文本":
                    _filename+=str(_rule[2])
                elif _rule[1]=="数字":
                    _filename+=str(int(_rule[2])+i)
                elif _rule[1]=="execl":
                    _sheet=self._execlData[str(_rule[2])]
                    _filename+=str(_sheet.cell(i+1,1).value)

            if _filename!="":
                self._savedoc._savelist.AddText(_filelink+"/"+_filename+_suffix)

        self._frame.Go(loadRunIndex)

    def __OpenFile(self):
        os.startfile(self._loadName._savelink.text())
        self.window()._info.ShowInfo(InfoType.SUCCESS,"提示:","保存文件夹已打开")
            

    def __Run(self):
        for i in range(len(self._savedoc._savelist._list)):
            if self._savedoc._savelist._list[i].isChecked():
                _doc=Document(self._docLink)
                self.__ReplaceData(_doc,i)
                _doc.save(self._savedoc._savelist._list[i].text())
                self._savedoc._savelist._list[i].AddTip()

        self.window()._info.ShowInfo(InfoType.SUCCESS,"提示:","运行完毕")

    def __ReplaceData(self,Doc,Column):
        _execly=1

        for paragraph in Doc.paragraphs:
            for run in paragraph.runs:
                if self._docKey in run.text:
                    _list=run.text.split(self._docKey)
       
                    run.text=""
        
                    for i in range(len(_list)-1):
                        print(_execly,Column+1)
                        run.text +=_list[i]+str(self._execlData.worksheets[self._execlDataIndex].cell(_execly,Column+1).value)
                        _execly+=1

                    run.text +=_list[-1]

class WaitWidget(QWidget):
    def __init__(self,*args, **kwargs):
        super(WaitWidget, self).__init__(*args, **kwargs)
        self.__SetLayout()

    def __SetLayout(self):
        _hbox=QHBoxLayout()
        _hbox.setContentsMargins(0,0,0,0)
        _hbox.setSpacing(0)

        self.setLayout(_hbox)

        _vbox=QVBoxLayout()
        _vbox.setContentsMargins(0,0,0,0)
        _vbox.setSpacing(15)

        self._progress=CircularProgressBar()
        self._progress.setFixedSize(200,200)
  
        self._text=LabelNormal()
        self._text.setWrapText(False)
        self._text.setAlignment(Qt.AlignCenter)
        _textfont=Font_(FontName="微软雅黑",FontSize=FontSize(Size=25,Shift=0,Min=10,Max=16))
        self._text.qss.SetFont(_textfont)
        self._text._qssplus.ApplyQss()
        
        _vbox.addStretch(1)
        _vbox.addWidget(self._progress,alignment=Qt.AlignHCenter)
        _vbox.addWidget(self._text,alignment=Qt.AlignHCenter)
        _vbox.addStretch(1)

        _hbox.addStretch(1)
        _hbox.addLayout(_vbox)
        _hbox.addStretch(1)

    def SetText(self,text:str):
        self._text.setText(text)
        self._text.adjustSize()

class LoadWord(QWidget):
    nextpressed=Signal(str)
    def __init__(self,*args, **kwargs):
        super(LoadWord, self).__init__(*args, **kwargs)
        self.setAcceptDrops(True) 
        self.__SetLayout()

    def __SetLayout(self):
        self.layouts=QVBoxLayout()
        self.layouts.setContentsMargins(40,40,40,40)
        self.layouts.setSpacing(10)
        self.layouts.addStretch(2)
        self.setLayout(self.layouts)

        self._word=IconView(Icon.word1,None,QSize(200,200))
        self._word.setFixedSize(QSize(200,200))
        _wordlayout=QHBoxLayout()
        _wordlayout.addStretch(1)
        _wordlayout.addWidget(self._word)
        _wordlayout.addStretch(1)
        self.layouts.addLayout(_wordlayout,1)

        _text=LabelNormal("拖拽Word文件或输入Word文件路径")
        _text.setWrapText(False)
        _text.setAlignment(Qt.AlignCenter)
        _textfont=Font_(FontName="微软雅黑",FontSize=FontSize(Size=25,Shift=0,Min=10,Max=16))
        _text.qss.SetFont(_textfont)
        _text._qssplus.ApplyQss()
        self.layouts.addWidget(_text,1)

        self._wordlink=LineEdit()
        self._wordlink.setPlaceholderText("请输入Word文件路径")
        self._wordlink.SetAddLink("请选择Word文档","","Word (*.doc *.docx);")
        self.layouts.addWidget(self._wordlink,1)

        self._next=ButtonNormal("下一步")
        self._next.setFixedSize(100,32)
        self._next.clicked.connect(self.__Next)

        _nextlayout=QHBoxLayout()
        _nextlayout.addStretch(1)
        _nextlayout.addWidget(self._next)
        _nextlayout.addStretch(1)

        self.layouts.addStretch(2)
        self.layouts.addLayout(_nextlayout,1)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event):
        for url in event.mimeData().urls():
            file_path = str(url.toLocalFile())
            # 处理文件路径，可以使用os模块进行进一步操作
            if self.IsWord(file_path):
                self._wordlink.setText(file_path)
            else:
                self.window()._info.ShowInfo(InfoType.WARNING,"提示:","拖入文件不是Word文件。")

    def __Next(self):
        self.nextpressed.emit(self._wordlink.text())

    def IsWord(self,link:str):
        if os.path.exists(link):
            _suffix=os.path.splitext(link)[1]
            if _suffix==".doc" or _suffix==".docx":
                return True
            else:
                return False
        else:
            return False

class LoadKey(QWidget):
    beforpressed=Signal(int)
    checkpressed=Signal(str)
    nextpressed=Signal(str)
    def __init__(self,*args, **kwargs):
        super(LoadKey, self).__init__(*args, **kwargs)
        self.__SetLayout()

    def __SetLayout(self):
        self.layouts=QVBoxLayout()
        self.layouts.setContentsMargins(40,40,40,40)
        self.layouts.setSpacing(10)
        self.layouts.addStretch(2)
        self.setLayout(self.layouts)

        _text=LabelNormal("设置Word的关键词并核对关键词数量是否和设置的数量一致")
        _text.setWrapText(False)
        _text.setAlignment(Qt.AlignCenter)
        _textfont=Font_(FontName="微软雅黑",FontSize=FontSize(Size=25,Shift=0,Min=10,Max=16))
        _text.qss.SetFont(_textfont)
        _text._qssplus.ApplyQss()

        self.layouts.addWidget(_text,1)

        _text1=LabelNormal("Word模板路径:")
        _text1.setFixedWidth(textwidth)

        self._wordlink=LineEdit()
        self._wordlink.setPlaceholderText("Word文件路径")
        self._wordlink.setReadOnly(True)

        _wordlayout=QHBoxLayout()
        _wordlayout.setContentsMargins(5,0,5,0)
        _wordlayout.setSpacing(5)
        _wordlayout.addWidget(_text1)
        _wordlayout.addWidget(self._wordlink)

        self.layouts.addLayout(_wordlayout,1)

        _text2=LabelNormal("Word模板关键词:")
        _text2.setFixedWidth(textwidth)

        self._key=LineEdit()
        self._key.setPlaceholderText("请设置Word模板中的关键词")

        _keylayout=QHBoxLayout()
        _keylayout.setContentsMargins(5,0,5,0)
        _keylayout.setSpacing(5)
        _keylayout.addWidget(_text2)
        _keylayout.addWidget(self._key)

        self.layouts.addLayout(_keylayout,1)

        self._before=ButtonNormal("上一步")
        self._before.setFixedSize(100,32)
        self._before.clicked.connect(self.__Before)

        self._checkKey=ButtonNormal("检查关键词")
        self._checkKey.setFixedSize(100,32)
        self._checkKey.clicked.connect(self.__Check)

        self._next=ButtonNormal("下一步")
        self._next.setFixedSize(100,32)
        self._next.clicked.connect(self.__Next)

        _nextlayout=QHBoxLayout()
        _nextlayout.addStretch(1)
        _nextlayout.addWidget(self._before)
        _nextlayout.addWidget(self._checkKey)
        _nextlayout.addWidget(self._next)
        _nextlayout.addStretch(1)

        self.layouts.addStretch(2)

        self.layouts.addLayout(_nextlayout,1)

    def __Next(self):
        self.nextpressed.emit(self._key.text())

    def __Before(self):
        self.beforpressed.emit(loadWordIndex)

    def __Check(self):
        if self._key.text()=="":
            self.window()._info.ShowInfo(InfoType.WARNING,"提示:","关键词不能为空。")
            return
        
        self.checkpressed.emit(self._key.text())

class LoadExecl(QWidget):
    beforpressed=Signal(int)
    nextpressed=Signal(str)
    def __init__(self,*args, **kwargs):
        super(LoadExecl, self).__init__(*args, **kwargs)
        self.setAcceptDrops(True) 
        self.__SetLayout()

    def __SetLayout(self):
        self.layouts=QVBoxLayout()
        self.layouts.setContentsMargins(40,40,40,40)
        self.layouts.setSpacing(10)
        self.layouts.addStretch(2)
        self.setLayout(self.layouts)

        self._execl=IconView(Icon.execl,None,QSize(200,200))
        self._execl.setFixedSize(QSize(200,200))
        _wordlayout=QHBoxLayout()
        _wordlayout.addStretch(1)
        _wordlayout.addWidget(self._execl)
        _wordlayout.addStretch(1)
        self.layouts.addLayout(_wordlayout,1)

        _text=LabelNormal("拖拽Execl文件或输入Execl文件路径")
        _text.setWrapText(False)
        _text.setAlignment(Qt.AlignCenter)
        _textfont=Font_(FontName="微软雅黑",FontSize=FontSize(Size=25,Shift=0,Min=10,Max=16))
        _text.qss.SetFont(_textfont)
        _text._qssplus.ApplyQss()
        self.layouts.addWidget(_text,1)

        self._execllink=LineEdit()
        self._execllink.setPlaceholderText("请输入Execl文件路径")
        self._execllink.SetAddLink("请选择Execl文档","","Execl (*.xls *.xlsx);")
        self.layouts.addWidget(self._execllink,1)

        self._before=ButtonNormal("上一步")
        self._before.setFixedSize(100,32)
        self._before.clicked.connect(self.__Before)

        self._next=ButtonNormal("下一步")
        self._next.setFixedSize(100,32)
        self._next.clicked.connect(self.__Next)

        _nextlayout=QHBoxLayout()
        _nextlayout.addStretch(1)
        _nextlayout.addWidget(self._before)
        _nextlayout.addWidget(self._next)
        _nextlayout.addStretch(1)

        self.layouts.addStretch(2)

        self.layouts.addLayout(_nextlayout,1)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event):
        for url in event.mimeData().urls():
            file_path = str(url.toLocalFile())
            # 处理文件路径，可以使用os模块进行进一步操作
            if self.IsExecl(file_path):
                self._execllink.setText(file_path)
            else:
                self.window()._info.ShowInfo(InfoType.WARNING,"提示:","拖入文件不是Execl文件。")

    def __Next(self):
        self.nextpressed.emit(self._execllink.text())

    def __Before(self):
        self.beforpressed.emit(loadKeyIndex)

    def IsExecl(self,link:str):
        if os.path.exists(link)==False:
            return False
        
        _suffix=os.path.splitext(link)[1]
        if _suffix==".xls" or _suffix==".xlsx":
            return True
        else:
            return False
  
class LoadData(QWidget):
    beforpressed=Signal(int)
    nextpressed=Signal(int)
    def __init__(self,*args, **kwargs):
        super(LoadData, self).__init__(*args, **kwargs)
        self.__SetLayout()

    def __SetLayout(self):
        self.layouts=QVBoxLayout()
        self.layouts.setContentsMargins(40,40,40,40)
        self.layouts.setSpacing(10)
        self.layouts.addStretch(2)
        self.setLayout(self.layouts)

        _text=LabelNormal("选择Execl数据表")
        _text.setWrapText(False)
        _text.setAlignment(Qt.AlignCenter)
        _textfont=Font_(FontName="微软雅黑",FontSize=FontSize(Size=25,Shift=0,Min=10,Max=16))
        _text.qss.SetFont(_textfont)
        _text._qssplus.ApplyQss()
        self.layouts.addWidget(_text,1)

        _text1=LabelNormal("Execl文件路径:")
        _text1.setFixedWidth(textwidth)

        self._execllink=LineEdit()
        self._execllink.setPlaceholderText("Eexcl文件路径")
        self._execllink.setReadOnly(True)

        _execllayout=QHBoxLayout()
        _execllayout.setContentsMargins(5,0,5,0)
        _execllayout.setSpacing(5)
        _execllayout.addWidget(_text1)
        _execllayout.addWidget(self._execllink)

        self.layouts.addLayout(_execllayout,1)

        _text2=LabelNormal("Execl数据表:")
        _text2.setFixedWidth(textwidth)

        self._execldatapage=ComBoBox()
        self._execldatapage.setFixedHeight(32)

        _execldatalayout=QHBoxLayout()
        _execldatalayout.setContentsMargins(5,0,5,0)
        _execldatalayout.setSpacing(5)
        _execldatalayout.addWidget(_text2)
        _execldatalayout.addWidget(self._execldatapage)

        self.layouts.addLayout(_execldatalayout,1)

        self._before=ButtonNormal("上一步")
        self._before.setFixedSize(100,32)
        self._before.clicked.connect(self.__Before)

        self._next=ButtonNormal("下一步")
        self._next.setFixedSize(100,32)
        self._next.clicked.connect(self.__Next)

        _nextlayout=QHBoxLayout()
        _nextlayout.addStretch(1)
        _nextlayout.addWidget(self._before)
        _nextlayout.addWidget(self._next)
        _nextlayout.addStretch(1)

        self.layouts.addStretch(2)

        self.layouts.addLayout(_nextlayout,1)

    def __Next(self):
        self.nextpressed.emit(self._execldatapage._listwidget._check)

    def __Before(self):
        self.beforpressed.emit(loadExeclIndex)

class LoadName(QWidget):
    beforpressed=Signal(int)
    nextpressed=Signal(list)
    def __init__(self,*args, **kwargs):
        super(LoadName, self).__init__(*args, **kwargs)
        self.__SetLayout()

    def __SetLayout(self):
        self.layouts=QVBoxLayout()
        self.layouts.setContentsMargins(40,40,40,40)
        self.layouts.setSpacing(10)
        self.layouts.addStretch(2)
        self.setLayout(self.layouts)

        _text=LabelNormal("设置保存信息")
        _text.setWrapText(False)
        _text.setAlignment(Qt.AlignCenter)
        _textfont=Font_(FontName="微软雅黑",FontSize=FontSize(Size=25,Shift=0,Min=10,Max=16))
        _text.qss.SetFont(_textfont)
        _text._qssplus.ApplyQss()
        self.layouts.addWidget(_text,1)

        _text1=LabelNormal("保存文件夹路径:")
        _text1.setFixedWidth(textwidth)

        self._savelink=LineEdit()
        self._savelink.setPlaceholderText("请输入保存文件夹路径")
        self._savelink.SetAddLink("请选择保存文件夹路径","","",True)
        
        _savelayout=QHBoxLayout()
        _savelayout.setContentsMargins(5,0,5,0)
        _savelayout.setSpacing(5)
        _savelayout.addWidget(_text1)
        _savelayout.addWidget(self._savelink)

        self.layouts.addLayout(_savelayout,1)

        _text2=LabelNormal("文档命名格式:")
        _text2.setFixedWidth(textwidth)

        self._nameRule=NameRuleList()
        self._nameRule.setFixedHeight(50)

        _namerulelayout=QHBoxLayout()
        _namerulelayout.setContentsMargins(5,0,5,0)
        _namerulelayout.setSpacing(5)
        _namerulelayout.addWidget(_text2)
        _namerulelayout.addWidget(self._nameRule)

        self.layouts.addLayout(_namerulelayout,1)

        self._before=ButtonNormal("上一步")
        self._before.setFixedSize(100,32)
        self._before.clicked.connect(self.__Before)

        self._next=ButtonNormal("下一步")
        self._next.setFixedSize(100,32)
        self._next.clicked.connect(self.__Next)

        _nextlayout=QHBoxLayout()
        _nextlayout.addStretch(1)
        _nextlayout.addWidget(self._before)
        _nextlayout.addWidget(self._next)
        _nextlayout.addStretch(1)

        self.layouts.addStretch(2)

        self.layouts.addLayout(_nextlayout,1)

    def __Next(self):
        self.nextpressed.emit(self._nameRule._rulelist)

    def __Before(self):
        self.beforpressed.emit(loadDataIndex)

class SaveDoc(QWidget):
    firsted=Signal(int)
    beforpressed=Signal(int)
    openfiled=Signal()
    runed=Signal()
    def __init__(self,*args, **kwargs):
        super(SaveDoc, self).__init__(*args, **kwargs)
        self.__SetLayout()

    def __SetLayout(self):
        self.layouts=QVBoxLayout()
        self.layouts.setContentsMargins(40,30,40,40)
        self.layouts.setSpacing(10)
        # self.layouts.addStretch(1)
        self.setLayout(self.layouts)

        _text=LabelNormal("预生成文档")
        _text.setFixedHeight(40)
        _text.setWrapText(False)
        _textfont=Font_(FontName="微软雅黑",FontSize=FontSize(Size=25,Shift=0,Min=10,Max=16))
        _text.qss.SetFont(_textfont)
        _text._qssplus.ApplyQss()
        self.layouts.addWidget(_text,1)

        self._savelist=TaskList()
        self._savelist.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.layouts.addWidget(self._savelist,1)

        self._before=ButtonNormal("上一步")
        self._selectAll=ButtonNormal("全选")
        self._selectTran=ButtonNormal("反选")
        self._selectRun=ButtonNormal("运行选中")
        self._openFile=ButtonNormal("打开文件夹")
        self._openFirst=ButtonNormal("重新载入Word")

        self._before.setFixedSize(80,32)
        self._selectAll.setFixedSize(60,32)
        self._selectTran.setFixedSize(60,32)
        self._selectRun.setFixedSize(90,32)
        self._openFile.setFixedSize(90,32)
        self._openFirst.setFixedSize(110,32)

        self._selectAll.clicked.connect(self._savelist.SelectAll)
        self._selectTran.clicked.connect(self._savelist.SelectTran)
        self._selectRun.clicked.connect(self.runed.emit)
        self._openFile.clicked.connect(self.openfiled.emit)
        self._openFirst.clicked.connect(self.__First)
        self._before.clicked.connect(self.__Before)

        _btnlayout=QHBoxLayout()
        _btnlayout.setContentsMargins(5,10,5,0)
        _btnlayout.setSpacing(5)
        _btnlayout.addStretch(1)
        _btnlayout.addWidget(self._before)
        _btnlayout.addWidget(self._selectAll)
        _btnlayout.addWidget(self._selectTran)
        _btnlayout.addWidget(self._selectRun)
        _btnlayout.addWidget(self._openFile)
        _btnlayout.addWidget(self._openFirst)
        _btnlayout.addStretch(1)

        # self.layouts.addStretch(1)

        self.layouts.addLayout(_btnlayout,1)

    def __Before(self):
        self.beforpressed.emit(loadNameIndex)

    def __First(self):
        self.firsted.emit(loadWordIndex)

class LoadExeclThread(QThread):
    loaded=Signal(bytes)
    errored=Signal()
    def __init__(self,ExeclLink:str):
        super(LoadExeclThread,self).__init__()
        self._execllink=ExeclLink
        
    def run(self):
        xls=None
        self.sleep(1)
        try:
            xls = load_workbook(self._execllink)
        except:
            self.errored.emit()
            return

        self.loaded.emit(xls)

class LoadWordThread(QThread):
    loaded=Signal(bytes)
    errored=Signal()
    def __init__(self,WordLink:str):
        super(LoadWordThread,self).__init__()
        self._wordlink=WordLink
        
    def run(self):
        self.sleep(1)

        doc=None

        try:
            doc = Document(self._wordlink)
        except:
            self.errored.emit()
            return

        self.loaded.emit(doc)

class KeyThread(QThread):
    checked=Signal(int)
    def __init__(self,Key:str,doc):
        super(KeyThread,self).__init__()
        self._doc=doc
        self._key=Key
        
    def run(self):
        self.sleep(1)

        _number=0

        for paragraph in self._doc.paragraphs:
            _number+=paragraph.text.count(self._key)

        self.checked.emit(_number)

if __name__ == '__main__':
   
    app =QApplication(sys.argv)
    w_=MainWidget()
    w_.show()
    sys.exit(app.exec())


